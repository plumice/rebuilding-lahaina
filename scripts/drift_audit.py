#!/usr/bin/env python3
"""
Per-paragraph drift audit: V31.docx vs thesis_source.txt.

For each drifted section (2A, 3B, 3C, 3D), align paragraphs between
the two sources using fuzzy match, then report:
  - Paragraphs ONLY in thesis_source.txt (added post-V31)
  - Paragraphs ONLY in V31 (removed from website)
  - Paragraphs in both but MODIFIED (with a unified diff)

Output: docs/drift-audit/2026-05-17-v31-vs-website-drift.md
"""
import re
import sys
import difflib
from pathlib import Path

CANONICAL_DOCX = "/Users/akhil/Library/CloudStorage/OneDrive-Personal/THESIS/CURRENT WRITING/Thesis_V32.docx"
CANONICAL_EXTRACT = "/tmp/thesis_v32_extracted.txt"
SRC_PATH = "thesis_source.txt"
OUT_PATH = "docs/drift-audit/2026-05-18-v32-vs-website-drift.md"

# Back-compat (legacy variable name still used elsewhere)
V31_PATH = CANONICAL_EXTRACT

SECTIONS_TO_AUDIT = ["2A", "3B", "3C", "3D"]
MATCH_THRESHOLD = 0.55  # similarity ratio to consider paragraphs "the same"
MIN_PARA_LEN = 40  # ignore very short lines (headings, single-word lines)


HEAD = re.compile(r"^([1-5][A-F])\s*/\s*(.+)$", re.MULTILINE)


def split_sections(text):
    matches = list(HEAD.finditer(text))
    out = {}
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        out[m.group(1)] = body
    return out


def split_paragraphs(text):
    """Split section text into non-trivial paragraphs.

    thesis_source.txt uses single-newline paragraphs; V31 extract may too.
    Treat each non-empty line as a paragraph candidate.
    """
    paras = []
    for line in text.split("\n"):
        line = line.strip()
        if len(line) >= MIN_PARA_LEN:
            paras.append(line)
    return paras


def normalize(s):
    """Loose normalization for fuzzy matching."""
    s = s.lower()
    s = re.sub(r"[—–-]+", " ", s)  # all dashes → space
    s = re.sub(r"[^\w\s]", " ", s)  # strip punct
    s = re.sub(r"\s+", " ", s).strip()
    return s


def similarity(a, b):
    return difflib.SequenceMatcher(None, normalize(a), normalize(b)).ratio()


def align(v_paras, s_paras):
    """Match each src paragraph to its closest V31 paragraph.

    Returns: (matched_pairs, src_only, v31_only)
      matched_pairs: list of (v_idx, s_idx, ratio, v_para, s_para)
      src_only: list of (s_idx, s_para)  — added post-V31
      v31_only: list of (v_idx, v_para)  — removed from website
    """
    pairs = []
    used_v = set()
    src_only = []

    # Greedy match: for each src paragraph, find best unmatched v31 paragraph
    for s_idx, s_para in enumerate(s_paras):
        best_v = -1
        best_ratio = 0.0
        for v_idx, v_para in enumerate(v_paras):
            if v_idx in used_v:
                continue
            r = similarity(s_para, v_para)
            if r > best_ratio:
                best_ratio = r
                best_v = v_idx
        if best_ratio >= MATCH_THRESHOLD:
            used_v.add(best_v)
            pairs.append((best_v, s_idx, best_ratio, v_paras[best_v], s_para))
        else:
            src_only.append((s_idx, s_para))

    v31_only = [(v_idx, v_para) for v_idx, v_para in enumerate(v_paras) if v_idx not in used_v]
    return pairs, src_only, v31_only


def render_diff(a, b, label_a="V31", label_b="website"):
    """Unified diff snippet for a modified pair."""
    diff = difflib.unified_diff(
        a.split(". "), b.split(". "),
        lineterm="", fromfile=label_a, tofile=label_b, n=1
    )
    return "\n".join(list(diff)[:30])


def short(s, n=200):
    return s if len(s) <= n else s[:n].rsplit(" ", 1)[0] + "…"


def ensure_canonical_extract():
    """Re-extract canonical .docx → .txt if extract missing or older than docx."""
    docx_p = Path(CANONICAL_DOCX)
    txt_p = Path(CANONICAL_EXTRACT)
    if not docx_p.exists():
        sys.stderr.write(f"!! Canonical thesis not found: {docx_p}\n")
        sys.exit(2)
    if txt_p.exists() and txt_p.stat().st_mtime >= docx_p.stat().st_mtime:
        return  # extract is fresh
    sys.stderr.write(f"Extracting {docx_p.name} → {txt_p} ...\n")
    from docx import Document
    doc = Document(str(docx_p))
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for i, t in enumerate(doc.tables):
        lines.append(f"\n[TABLE {i+1}]")
        for row in t.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    txt_p.write_text("\n".join(lines), encoding="utf-8")


def main():
    root = Path(__file__).resolve().parent.parent
    ensure_canonical_extract()
    v31_text = Path(CANONICAL_EXTRACT).read_text(encoding="utf-8", errors="replace")
    src_text = (root / SRC_PATH).read_text(encoding="utf-8", errors="replace")

    v31_sec = split_sections(v31_text)
    src_sec = split_sections(src_text)

    out_dir = root / "docs" / "drift-audit"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = root / OUT_PATH

    lines = []
    lines.append("# V31.docx vs Website Drift Audit")
    lines.append("")
    lines.append("**Generated:** 2026-05-17")
    lines.append(f"**Tool:** `scripts/drift_audit.py`")
    lines.append(f"**Match threshold:** {MATCH_THRESHOLD} (paragraphs with similarity below this are flagged as added/removed)")
    lines.append("")
    lines.append("Compares `/Users/akhil/Library/CloudStorage/OneDrive-Personal/THESIS/CURRENT WRITING/Thesis_V31.docx`")
    lines.append("against `thesis_source.txt` for each of the four drifted sections (2A, 3B, 3C, 3D).")
    lines.append("")
    lines.append("**For each drifted paragraph, decide: KEEP (it's intentional post-V31 author work) or RESTORE (remove and match V31).**")
    lines.append("")
    lines.append("---")
    lines.append("")

    summary = []

    for key in SECTIONS_TO_AUDIT:
        v_body = v31_sec.get(key, "")
        s_body = src_sec.get(key, "")
        v_paras = split_paragraphs(v_body)
        s_paras = split_paragraphs(s_body)
        pairs, src_only, v31_only = align(v_paras, s_paras)

        # Count modified (matched but with notable differences)
        modified = [p for p in pairs if p[2] < 0.95]

        summary.append((key, len(src_only), len(v31_only), len(modified)))

        lines.append(f"## Section {key}")
        lines.append("")
        lines.append(f"- V31 paragraphs: {len(v_paras)}")
        lines.append(f"- Website paragraphs: {len(s_paras)}")
        lines.append(f"- **Added to website (not in V31): {len(src_only)}**")
        lines.append(f"- **Removed from website (in V31 only): {len(v31_only)}**")
        lines.append(f"- **Modified (in both, but with differences): {len(modified)}**")
        lines.append("")

        if src_only:
            lines.append(f"### ⊕ Added to website ({len(src_only)})")
            lines.append("")
            for i, (s_idx, s_para) in enumerate(src_only, 1):
                lines.append(f"**A{i}** _(website paragraph #{s_idx + 1})_")
                lines.append("")
                lines.append(f"> {short(s_para, 600)}")
                lines.append("")
                lines.append("- [ ] KEEP (intentional post-V31 edit)")
                lines.append("- [ ] RESTORE (remove from website)")
                lines.append("")

        if v31_only:
            lines.append(f"### ⊖ Removed from website ({len(v31_only)})")
            lines.append("")
            for i, (v_idx, v_para) in enumerate(v31_only, 1):
                lines.append(f"**R{i}** _(V31 paragraph #{v_idx + 1})_")
                lines.append("")
                lines.append(f"> {short(v_para, 600)}")
                lines.append("")
                lines.append("- [ ] KEEP REMOVAL (website intentionally drops this)")
                lines.append("- [ ] RESTORE (add back from V31)")
                lines.append("")

        if modified:
            lines.append(f"### ✎ Modified ({len(modified)})")
            lines.append("")
            for i, (v_idx, s_idx, ratio, v_para, s_para) in enumerate(modified, 1):
                lines.append(f"**M{i}** _(similarity {ratio:.2f}; V31 #{v_idx + 1} → website #{s_idx + 1})_")
                lines.append("")
                lines.append(f"**V31:** {short(v_para, 500)}")
                lines.append("")
                lines.append(f"**Website:** {short(s_para, 500)}")
                lines.append("")
                lines.append("- [ ] KEEP (website improvement is intentional)")
                lines.append("- [ ] RESTORE (revert website to V31 wording)")
                lines.append("")

        lines.append("---")
        lines.append("")

    # Summary table at top
    summary_lines = ["## Summary", "", "| Section | Added | Removed | Modified |", "|---|---:|---:|---:|"]
    total_a = total_r = total_m = 0
    for key, a, r, m in summary:
        summary_lines.append(f"| {key} | {a} | {r} | {m} |")
        total_a += a; total_r += r; total_m += m
    summary_lines.append(f"| **TOTAL** | **{total_a}** | **{total_r}** | **{total_m}** |")
    summary_lines.append("")
    summary_lines.append("---")
    summary_lines.append("")

    # Insert summary after the intro section
    insert_at = next(i for i, l in enumerate(lines) if l.startswith("## Section "))
    final = lines[:insert_at] + summary_lines + lines[insert_at:]

    out_path.write_text("\n".join(final), encoding="utf-8")
    total_changes = total_a + total_r + total_m
    print(f"Wrote {out_path}")
    print(f"  Sections audited: {len(SECTIONS_TO_AUDIT)}")
    print(f"  Total drift items: {total_changes}")
    print(f"    added:    {total_a}")
    print(f"    removed:  {total_r}")
    print(f"    modified: {total_m}")


if __name__ == "__main__":
    main()
